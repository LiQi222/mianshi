import os
import sqlite3
import re
import io
from flask import Flask, request, jsonify, send_from_directory, send_file
from flask_cors import CORS
from flask_login import LoginManager, UserMixin, login_user, logout_user, login_required, current_user
from werkzeug.security import generate_password_hash, check_password_hash
import requests
import PyPDF2 as pdf
from docx import Document

# --- App & DB Setup ---
app = Flask(__name__, static_folder='static')
app.config['SECRET_KEY'] = os.urandom(24)
CORS(app)

DATABASE = 'database.db'

def get_db():
    db = sqlite3.connect(DATABASE)
    db.row_factory = sqlite3.Row
    return db

def init_db():
    with app.app_context():
        db = get_db()
        if os.path.exists('schema.sql'):
            with open('schema.sql', 'r') as f:
                db.executescript(f.read())
            db.commit()
            print("Database initialized.")

# --- User Authentication Setup ---
login_manager = LoginManager()
login_manager.init_app(app)

class User(UserMixin):
    def __init__(self, id, username):
        self.id = id
        self.username = username

@login_manager.user_loader
def load_user(user_id):
    db = get_db()
    user_row = db.execute('SELECT * FROM users WHERE id = ?', (user_id,)).fetchone()
    db.close()
    if user_row:
        return User(id=user_row['id'], username=user_row['username'])
    return None

# --- API Key Setup ---
try:
    api_key = os.getenv("DEEPSEEK_API_KEY")
    if not api_key:
        raise ValueError("未找到 DEEPSEEK_API_KEY 环境变量")
except Exception as e:
    print(f"API密钥配置错误: {e}")

# --- Frontend & Static Routes ---
@app.route('/')
def serve_index():
    return send_from_directory(app.static_folder, 'index.html')

@app.route('/agreement')
def agreement_page():
    return send_from_directory(app.static_folder, 'agreement.html')

@app.route('/privacy')
def privacy_page():
    return send_from_directory(app.static_folder, 'privacy.html')

# --- API Routes ---
@app.route('/api/register', methods=['POST'])
def register():
    data = request.get_json()
    username = data.get('username')
    password = data.get('password')

    if not re.match(r'^[a-zA-Z0-9]{5,10}$', username):
        return jsonify({"error": "用户名必须是5-10位的英文字母或数字"}), 400
    if not re.match(r'^(?=.*[a-zA-Z])(?=.*\d)[a-zA-Z\d]{6,18}$', password):
        return jsonify({"error": "密码必须是6-18位，且同时包含英文和数字"}), 400

    db = get_db()
    if db.execute('SELECT id FROM users WHERE username = ?', (username,)).fetchone():
        db.close()
        return jsonify({"error": "用户名已存在"}), 409

    hashed_password = generate_password_hash(password)
    db.execute('INSERT INTO users (username, password) VALUES (?, ?)', (username, hashed_password))
    db.commit()
    db.close()
    return jsonify({"success": "注册成功"}), 201

@app.route('/api/login', methods=['POST'])
def login():
    data = request.get_json()
    username = data.get('username')
    password = data.get('password')
    db = get_db()
    user_row = db.execute('SELECT * FROM users WHERE username = ?', (username,)).fetchone()
    db.close()
    if user_row and check_password_hash(user_row['password'], password):
        user = User(id=user_row['id'], username=user_row['username'])
        login_user(user)
        return jsonify({"success": "登录成功", "username": user.username})
    return jsonify({"error": "用户名或密码错误"}), 401

@app.route('/api/logout', methods=['POST'])
@login_required
def logout():
    logout_user()
    return jsonify({"success": "已退出登录"})

@app.route('/api/check_auth')
def check_auth():
    if current_user.is_authenticated:
        return jsonify({"is_authenticated": True, "username": current_user.username})
    return jsonify({"is_authenticated": False})

@app.route('/api/history')
@login_required # 只有登录用户可以获取历史
def get_history():
    db = get_db()
    history_rows = db.execute(
        'SELECT questions, timestamp FROM history WHERE user_id = ? ORDER BY timestamp DESC', 
        (current_user.id,)
    ).fetchall()
    db.close()
    history = [{"questions": row['questions'], "timestamp": row['timestamp']} for row in history_rows]
    return jsonify(history)

# --- 核心功能路由更新 ---
@app.route('/api/analyze', methods=['POST'])
def analyze_resume(): # 不再需要 @login_required
    try:
        if 'resume' not in request.files: return jsonify({"error": "没有找到简历文件"}), 400
        resume_file = request.files['resume']
        job_description = request.form.get('jd', '')
        if resume_file.filename == '': return jsonify({"error": "未选择文件"}), 400
        
        resume_text = extract_text_from_pdf(resume_file.stream)
        if "读取PDF时出错" in resume_text: return jsonify({"error": resume_text}), 500
        
        generated_questions = get_deepseek_response(resume_text, job_description)
        if "调用AI模型时出错" in generated_questions or "解析AI模型返回的数据时出错" in generated_questions:
             return jsonify({"error": generated_questions}), 500

        # 如果用户已登录，则保存历史记录
        if current_user.is_authenticated:
            db = get_db()
            db.execute('INSERT INTO history (user_id, questions) VALUES (?, ?)', (current_user.id, generated_questions))
            db.commit()
            db.close()

        return jsonify({"questions": generated_questions})
    except Exception as e:
        print(f"在 /analyze 端点发生严重错误: {e}")
        return jsonify({"error": "服务器内部发生严重错误"}), 500

# --- 新增：下载Word文档路由 ---
@app.route('/api/download_word', methods=['POST'])
@login_required # 只有登录用户可以下载
def download_word():
    data = request.get_json()
    content = data.get('content', '')
    
    document = Document()
    document.add_heading('AI生成的面试问题', 0)
    
    for line in content.split('\n'):
        if line.startswith('### '): document.add_heading(line.replace('### ', ''), level=3)
        elif line.startswith('## '): document.add_heading(line.replace('## ', ''), level=2)
        elif line.startswith('# '): document.add_heading(line.replace('# ', ''), level=1)
        elif line.strip().startswith('* '): document.add_paragraph(line.replace('* ', '').strip(), style='List Bullet')
        elif line.strip(): document.add_paragraph(line)
    
    f = io.BytesIO()
    document.save(f)
    f.seek(0)
    
    return send_file(f, as_attachment=True, download_name='面试问题.docx', mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document')

# --- Helper Functions ---
def get_deepseek_response(resume_text, job_description):
    url = "https://api.deepseek.com/chat/completions"
    prompt_content = f"请根据以下简历和职位描述，生成面试问题。\n\n职位描述:\n{job_description or '未提供'}\n\n简历文本:\n{resume_text}"
    payload = {"model": "deepseek-chat", "messages": [{"role": "user", "content": prompt_content}]}
    headers = {"Content-Type": "application/json", "Authorization": f"Bearer {api_key}"}
    try:
        response = requests.post(url, headers=headers, json=payload, timeout=100)
        response.raise_for_status()
        return response.json()['choices'][0]['message']['content']
    except Exception as e:
        return f"调用AI模型时出错: {str(e)}"

def extract_text_from_pdf(pdf_file):
    try:
        text = ""
        pdf_reader = pdf.PdfReader(pdf_file)
        for page in pdf_reader.pages:
            text += page.extract_text() or ""
        return text
    except Exception:
        return "读取PDF时出错"

# --- Main Execution ---
init_db()

if __name__ == '__main__':
    app.run(host='0.0.0.0', port=5000, debug=True)
